import docx

from core.schemas import LiteratureItem, TopicCard, WritingAsset


def _asset_map(items: list[WritingAsset]) -> dict[str, WritingAsset]:
    return {item.asset_type: item for item in items}


def _related_polish_assets(asset_type: str, polish_assets: list[WritingAsset]) -> list[WritingAsset]:
    return [item for item in polish_assets if asset_type in item.source_refs]


def export_assets_to_docx(
    topic_card: TopicCard | None,
    literature_items: list[LiteratureItem],
    writing_assets: list[WritingAsset],
    polish_assets: list[WritingAsset],
    output_path: str,
    template_label: str,
    section_template: list[tuple[str, str]],
) -> str:
    writing_asset_map = _asset_map(writing_assets)

    document = docx.Document()
    document.add_heading(template_label, level=0)

    document.add_heading("选题卡", level=1)
    if topic_card:
        document.add_paragraph(f"题目：{topic_card.title}")
        document.add_paragraph(f"研究问题：{topic_card.research_problem}")
        document.add_paragraph(f"研究对象：{topic_card.target_population}")
        document.add_paragraph(f"研究场景：{topic_card.context}")
        document.add_paragraph(f"关键词：{'、'.join(topic_card.keywords)}")
        document.add_paragraph(f"推荐方法：{'、'.join(topic_card.recommended_methods)}")
        document.add_paragraph(f"导师分析：{topic_card.mentor_analysis}")
    else:
        document.add_paragraph("暂无选题卡")

    document.add_heading("文献速读摘要", level=1)
    if literature_items:
        for item in literature_items:
            document.add_heading(item.file_name, level=2)
            document.add_paragraph(f"标题：{item.title}")
            document.add_paragraph(f"作者：{', '.join(item.authors)}")
            document.add_paragraph(f"年份：{item.year}")
            document.add_paragraph(f"摘要：{item.abstract}")
            document.add_paragraph(f"方法：{item.method}")
            document.add_paragraph(f"结论：{item.findings}")
    else:
        document.add_paragraph("暂无文献速读内容")

    document.add_heading("写作资产", level=1)
    document.add_paragraph(f"当前导出模板：{template_label}")
    document.add_heading(f"{template_label}正文结构", level=1)
    for asset_type, section_title in section_template:
        document.add_heading(section_title, level=2)
        asset = writing_asset_map.get(asset_type)
        if asset:
            document.add_paragraph(asset.content)
            related = _related_polish_assets(asset_type, polish_assets)
            for result in related:
                document.add_heading(result.title, level=3)
                document.add_paragraph(result.content)
        else:
            document.add_paragraph("暂无该部分内容")

    extras = [
        asset for asset in writing_assets if asset.asset_type not in {asset_type for asset_type, _ in section_template}
    ]
    if extras:
        document.add_heading("其他写作资产", level=1)
        for asset in extras:
            document.add_heading(asset.title or asset.asset_type, level=2)
            document.add_paragraph(asset.content)

    document.save(output_path)
    return output_path
